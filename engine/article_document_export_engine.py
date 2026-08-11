# -*- coding: utf-8 -*-

from pathlib import Path
from datetime import datetime, timezone
import json
import re
from engine.article_html_sanitization_engine import ArticleHTMLSanitizationEngine


class ArticleDocumentExportEngine:
    """
    Final Real PDF + DOCX Export Engine v2.0
    """

    VERSION = "2.0"

    def __init__(self):
        self.output = Path("output/documents")
        self.output.mkdir(parents=True, exist_ok=True)
        self.html_sanitizer = ArticleHTMLSanitizationEngine()

    def slugify(self, text):
        text = re.sub(r"[^\w\s-]", "", text)
        return text.strip().replace(" ", "_")

    def create_pdf(self, path, text):

        from reportlab.platypus import SimpleDocTemplate, Paragraph
        from reportlab.lib.styles import getSampleStyleSheet

        doc = SimpleDocTemplate(str(path))

        styles = getSampleStyleSheet()

        story = [Paragraph(text, styles["Normal"])]

        doc.build(story)

    def create_docx(self, path, text):

        from docx import Document

        doc = Document()

        doc.add_heading("AR-Vet Article", level=1)

        doc.add_paragraph(text)

        doc.save(str(path))

    def build(self, article):

        title = article.get("title", "article")

        slug = self.slugify(title)

        folder = self.output / slug

        folder.mkdir(parents=True, exist_ok=True)

        text = self.html_sanitizer.process(article.get("content", title)).get(
            "html", article.get("content", title)
        )

        pdf_path = folder / (slug + ".pdf")

        docx_path = folder / (slug + ".docx")

        self.create_pdf(pdf_path, text)

        self.create_docx(docx_path, text)

        manifest = {
            "title": title,
            "created": datetime.now(timezone.utc).isoformat(),
            "pdf": str(pdf_path),
            "docx": str(docx_path),
            "ready": True,
        }

        manifest_path = folder / "export_manifest.json"

        manifest_path.write_text(
            json.dumps(manifest, ensure_ascii=False, indent=4), encoding="utf-8"
        )

        return {
            "title": title,
            "formats": {"pdf": True, "docx": True},
            "pdf_path": str(pdf_path),
            "docx_path": str(docx_path),
            "manifest": str(manifest_path),
            "ready": True,
        }

    def info(self):

        return {
            "engine": "Real Document Export Engine",
            "version": self.VERSION,
            "status": "production",
            "pdf": True,
            "docx": True,
            "real_files": True,
        }
